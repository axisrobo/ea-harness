"""
from_document.py — Extract architecture requirements from text documents.

Supports: .pdf  .docx  .md  .txt

Strategy: Extract text from the file, then use Claude API to identify
architecture-relevant information and map it to the req.yaml schema.

Usage:
    python from_document.py -i requirements.pdf -o partial-req.yaml
    python from_document.py -i design-doc.docx -o partial-req.yaml
    ANTHROPIC_API_KEY=... python from_document.py -i brief.txt -o partial-req.yaml
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

import yaml

sys.path.insert(0, str(Path(__file__).parent))
from normalizer import (
    PartialReq, PartialApplication, PartialComponent, PartialInteraction,
    PartialUserAuth, Confidence, fv, partial_req_to_yaml
)


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_text_pdf(path: str) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages)
    except ImportError:
        # Fallback: try pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except ImportError:
            return ""


def _extract_text_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except ImportError:
        return ""


def extract_text(path: str) -> str:
    """Extract plain text from a document file."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        text = _extract_text_pdf(path)
    elif ext in (".docx", ".doc"):
        text = _extract_text_docx(path)
    elif ext in (".md", ".txt", ".rst"):
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()
    else:
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()
    return text


# ── LLM extraction ────────────────────────────────────────────────────────────

EXTRACTION_PROMPT = """You are an enterprise architecture analyst extracting structured requirements from a document.

Read the following document and extract all architecture-relevant information. 
Focus on: physical deployment locations, application names, technical stack, 
integration points, authentication mechanisms, data sensitivity, and security requirements.

Document content:
---
{document_text}
---

Extract the information into this exact JSON structure. Use null for unknown fields. 
Be conservative: only extract what is explicitly stated, do not infer.
Output ONLY the JSON, no explanation:

{{
  "project": {{
    "name": null,
    "id": null,
    "scope": null,
    "department": null,
    "business_purpose": null
  }},
  "applications": [
    {{
      "name": "app name",
      "type": "new|existing|modified|unknown",
      "owner": "org_it|biz_owned|third_party|unknown",
      "vendor": null,
      "dc_or_region": "physical location if stated",
      "country": "CN|US|DE|etc if stated",
      "platform": "private_dc|aws|azure|saas|unknown",
      "zone_subnet": "DMZ|App Zone|DB Zone|VPC|Subnet|unknown",
      "infra_owner": "InfraSec|BizIT|department name|unknown"
    }}
  ],
  "components": [
    {{
      "app_id": "which app this belongs to",
      "name": "component name",
      "type": "FE|BE|DB|MQ|IP|LB|SEC",
      "language": null,
      "framework": null,
      "runtime": null,
      "sensitivity": "Company Restricted|Company Confidential|Company Internal|null"
    }}
  ],
  "interactions": [
    {{
      "from": "initiator component name",
      "to": "provider component name",
      "protocol": "HTTPS|Kafka|SFTP|JDBC|RFC|TCP|null",
      "port": null,
      "auth_method": "OAuth2.0|mTLS|SAML|SASL/SCRAM|PWD|null"
    }}
  ],
  "user_auth": [
    {{
      "entry_point": "which application",
      "user_roles": ["role1", "role2"],
      "auth_server": "ADFS|EnterpriseID|EntraID|null",
      "auth_protocol": "SAML|CAS|OAuth2_AuthCode|OIDC|null",
      "authorization": "RBAC|ABAC|PBAC|DAC|null"
    }}
  ],
  "credentials": [
    {{
      "environment": "azure|aws|private_dc",
      "solution": "Azure Key Vault|AWS Secrets Manager|K8s Secrets|null"
    }}
  ],
  "data_encryption": [
    {{
      "component": "DB or storage name",
      "at_rest": true,
      "at_rest_method": "AES-256|TDE|null",
      "cross_border": false,
      "cross_border_compliance": "GDPR|中国数据安全法|null"
    }}
  ],
  "constraints": ["list of explicit constraints mentioned"],
  "gaps_noted": ["information that seems important but is missing from the document"]
}}
"""


def call_llm_extraction(document_text: str, source: str) -> dict:
    """Call Claude API to extract structured requirements from document text."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return {"error": "ANTHROPIC_API_KEY not set"}

    import urllib.request

    # Truncate very long documents to stay within context limits
    max_chars = 80000
    if len(document_text) > max_chars:
        document_text = document_text[:max_chars] + "\n... [document truncated]"

    prompt = EXTRACTION_PROMPT.format(document_text=document_text)

    payload = json.dumps({
        "model": "claude-sonnet-4-6",
        "max_tokens": 4096,
        "messages": [{"role": "user", "content": prompt}]
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        method="POST"
    )

    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            response = json.loads(resp.read().decode("utf-8"))
    except Exception as e:
        return {"error": str(e)}

    text = ""
    for block in response.get("content", []):
        if block.get("type") == "text":
            text = block["text"]
            break

    json_match = re.search(r'\{[\s\S]+\}', text)
    if not json_match:
        return {"error": "No JSON found in LLM response"}

    try:
        return json.loads(json_match.group())
    except json.JSONDecodeError as e:
        return {"error": f"JSON parse error: {e}"}


# ── Map LLM output → PartialReq ───────────────────────────────────────────────

def map_to_partial_req(extracted: dict, source_file: str) -> PartialReq:
    req = PartialReq(source_tool="arch-req-from-doc", source_file=source_file)
    SRC = f"document:{Path(source_file).name}"
    CONF = Confidence.MEDIUM   # document extraction is medium confidence

    if "error" in extracted:
        req.gaps.append(f"LLM extraction error: {extracted['error']}")
        return req

    # Project
    proj = extracted.get("project", {})
    if proj.get("name"):
        req.project_name = fv(proj["name"], CONF, SRC)
    if proj.get("id"):
        req.project_id = fv(proj["id"], CONF, SRC)
    if proj.get("scope"):
        req.project_scope = fv(proj["scope"], CONF, SRC)
    if proj.get("department"):
        req.department = fv(proj["department"], CONF, SRC)

    # Applications
    for i, app_data in enumerate(extracted.get("applications", [])):
        app = PartialApplication(id=f"doc_app_{i+1}")
        if app_data.get("name"):
            app.name = fv(app_data["name"], CONF, SRC)
        if app_data.get("type") and app_data["type"] != "unknown":
            app.type = fv(app_data["type"], CONF, SRC)
        if app_data.get("owner") and app_data["owner"] != "unknown":
            app.owner = fv(app_data["owner"], CONF, SRC)
        if app_data.get("dc_or_region"):
            app.dc_or_region = fv(app_data["dc_or_region"], CONF, SRC)
        if app_data.get("country"):
            app.country = fv(app_data["country"], CONF, SRC)
        if app_data.get("platform") and app_data["platform"] != "unknown":
            app.platform = fv(app_data["platform"], CONF, SRC)
        if app_data.get("zone_subnet") and app_data["zone_subnet"] != "unknown":
            app.zone_subnet = fv(app_data["zone_subnet"], CONF, SRC)
        if app_data.get("infra_owner") and app_data["infra_owner"] != "unknown":
            app.infra_owner = fv(app_data["infra_owner"], CONF, SRC)
        req.applications.append(app)

    # Components
    for i, comp_data in enumerate(extracted.get("components", [])):
        comp = PartialComponent(
            id=f"doc_comp_{i+1}",
            app_id=comp_data.get("app_id", "unknown")
        )
        if comp_data.get("name"):
            comp.name = fv(comp_data["name"], CONF, SRC)
        if comp_data.get("type"):
            comp.comp_type = fv(comp_data["type"], CONF, SRC)
        for field in ("language", "framework", "runtime", "sensitivity"):
            val = comp_data.get(field)
            if val:
                setattr(comp, field, fv(val, CONF, SRC))
        req.components.append(comp)

    # Interactions
    for i, iact in enumerate(extracted.get("interactions", [])):
        interaction = PartialInteraction(id=f"doc_int_{i+1}")
        if iact.get("from"):
            interaction.from_component = fv(iact["from"], CONF, SRC)
        if iact.get("to"):
            interaction.to_component = fv(iact["to"], CONF, SRC)
        if iact.get("protocol"):
            interaction.protocol = fv(iact["protocol"], CONF, SRC)
        if iact.get("auth_method"):
            interaction.auth_method = fv(iact["auth_method"], Confidence.LOW, SRC,
                                         "From document — verify each interaction has auth")
        if iact.get("port"):
            interaction.port = fv(str(iact["port"]), CONF, SRC)
        req.interactions.append(interaction)

    # User auth
    for i, ua_data in enumerate(extracted.get("user_auth", [])):
        pua = PartialUserAuth(entry_point=ua_data.get("entry_point", f"entry_{i+1}"))
        if ua_data.get("user_roles"):
            pua.user_roles = fv(ua_data["user_roles"], CONF, SRC)
        if ua_data.get("auth_server"):
            pua.auth_server = fv(ua_data["auth_server"], CONF, SRC)
        if ua_data.get("auth_protocol"):
            pua.auth_protocol = fv(ua_data["auth_protocol"], CONF, SRC)
        if ua_data.get("authorization"):
            pua.authorization = fv(ua_data["authorization"], CONF, SRC)
        req.user_auth.append(pua)

    # Credentials
    for cred in extracted.get("credentials", []):
        if cred.get("solution"):
            req.credentials.append({**cred, "_source": SRC, "_confidence": "medium"})

    # Data encryption
    for enc in extracted.get("data_encryption", []):
        req.data_encryption.append({**enc, "_source": SRC, "_confidence": "medium"})

    # Constraints
    req.constraints = extracted.get("constraints", [])

    # Gaps noted by LLM
    for gap in extracted.get("gaps_noted", []):
        req.gaps.append(f"Document gap: {gap}")

    req.gaps.append(f"Document extraction is MEDIUM confidence — verify all values, "
                    f"especially authentication mechanisms and physical locations")
    return req


# ── Main ─────────────────────────────────────────────────────────────────────

def parse_document(input_path: str) -> PartialReq:
    text = extract_text(input_path)
    if not text.strip():
        req = PartialReq(source_tool="arch-req-from-doc", source_file=input_path)
        req.gaps.append(f"Could not extract text from {input_path}. "
                        f"Install pdfplumber (PDF) or python-docx (DOCX).")
        return req
    extracted = call_llm_extraction(text, input_path)
    return map_to_partial_req(extracted, input_path)


def main():
    parser = argparse.ArgumentParser(description="Extract requirements from document")
    parser.add_argument("-i", "--input", required=True, help="Input document file")
    parser.add_argument("-o", "--output", default=None, help="Output partial-req YAML")
    args = parser.parse_args()

    req = parse_document(args.input)
    out = partial_req_to_yaml(req)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(out)
        print(f"✓ Partial requirements written: {args.output}")
        print(f"  Applications found: {len(req.applications)}")
        print(f"  Components found: {len(req.components)}")
        print(f"  Interactions found: {len(req.interactions)}")
        if req.gaps:
            print("  Notes:")
            for g in req.gaps[:5]:
                print(f"    • {g}")
    else:
        print(out)


if __name__ == "__main__":
    main()
